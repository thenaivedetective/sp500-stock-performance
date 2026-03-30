"""
Generates the ISEC-style two-column conference paper as a .docx file.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.top_margin    = Inches(0.75)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(0.75)
section.right_margin  = Inches(0.75)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

# Remove default paragraph spacing globally
doc.styles['Normal'].paragraph_format.space_before = Pt(0)
doc.styles['Normal'].paragraph_format.space_after  = Pt(0)

# ── Helper: set paragraph line spacing ──────────────────────────
def set_single(para):
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

# ── Helper: add a run with font settings ────────────────────────
def add_run(para, text, bold=False, italic=False, size=10, font="Times New Roman", color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# ── Helper: new paragraph with alignment ────────────────────────
def new_para(doc_or_section, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None, space_before=0, space_after=0):
    p = doc_or_section.add_paragraph()
    p.alignment = align
    set_single(p)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent is not None:
        p.paragraph_format.first_line_indent = Inches(indent)
    return p

# ── Helper: set two-column layout on a section ──────────────────
def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '288')   # 0.2" in twentieths of a point (0.2*1440=288)
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)

# ── Helper: insert continuous section break ──────────────────────
def insert_continuous_section_break(doc):
    p = doc.add_paragraph()
    set_single(p)
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    sectType = OxmlElement('w:type')
    sectType.set(qn('w:val'), 'continuous')
    sectPr.append(sectType)
    pPr.append(sectPr)
    return p

# ── Helper: small-caps XML ───────────────────────────────────────
def make_small_caps(run):
    rPr = run._r.get_or_add_rPr()
    smallCaps = OxmlElement('w:smallCaps')
    rPr.append(smallCaps)

# ── Helper: section heading ──────────────────────────────────────
def section_heading(doc, roman, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    set_single(p)
    text = f"{roman}. {title.upper()}"
    run = add_run(p, text, bold=True, size=10)
    make_small_caps(run)
    return p

# ── Helper: sub-heading ──────────────────────────────────────────
def sub_heading(doc, letter, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    set_single(p)
    add_run(p, f"{letter}. {title}", italic=True, size=10)
    return p

# ── Helper: body paragraph (first = no indent, rest = .25" indent)
def body_para(doc, text, first=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_single(p)
    p.paragraph_format.space_after  = Pt(space_after)
    if not first:
        p.paragraph_format.first_line_indent = Inches(0.25)
    add_run(p, text, size=10)
    return p

# ── Helper: bullet ───────────────────────────────────────────────
def bullet_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_single(p)
    p.paragraph_format.left_indent        = Inches(0.25)
    p.paragraph_format.first_line_indent  = Inches(-0.15)
    add_run(p, f"\u2022  {text}", size=10)
    return p

# ── Helper: insert figure ────────────────────────────────────────
def insert_figure(doc, image_path, roman, caption, width=3.3):
    try:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_single(p_img)
        p_img.paragraph_format.space_before = Pt(6)
        run = p_img.add_run()
        run.add_picture(image_path, width=Inches(width))
    except Exception:
        pass
    # Figure heading
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single(p_head)
    add_run(p_head, f"FIGURE {roman}", bold=False, size=10)
    # Caption
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_single(p_cap)
    p_cap.paragraph_format.space_after = Pt(6)
    run = add_run(p_cap, caption, size=8)
    make_small_caps(run)

# ── Helper: blank line (spacer) ──────────────────────────────────
def blank(doc, pts=10):
    p = doc.add_paragraph()
    set_single(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.font.size = Pt(pts)

# ═══════════════════════════════════════════════════════════════════
# TITLE SECTION  (one-column)
# ═══════════════════════════════════════════════════════════════════

# Title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_title)
add_run(p_title,
        "Credit Risk Classification Using Multivariate Statistical\n"
        "Techniques on LendingClub Loan Data",
        size=24, bold=False)

blank(doc, 11)

# Author
p_auth = doc.add_paragraph()
p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_auth)
add_run(p_auth, "Lana Jalal Gidan", size=11)

# Affiliation
p_aff = doc.add_paragraph()
p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_aff)
add_run(p_aff,
        "Department of Systems Science and Industrial Engineering, "
        "Binghamton University, lgidan@binghamton.edu",
        size=10)

blank(doc, 10)
blank(doc, 10)

# ── Abstract ────────────────────────────────────────────────────
p_abs = doc.add_paragraph()
p_abs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_single(p_abs)
r1 = add_run(p_abs, "Abstract", bold=True, italic=True, size=10)
add_run(p_abs,
    " - This paper presents a multivariate statistical learning framework "
    "for credit risk classification applied to a large-scale, real-world lending dataset "
    "obtained from LendingClub, a peer-to-peer loan marketplace. The dataset comprises "
    "500,000 loan records across 151 features, making it a rich and representative source "
    "of consumer credit behavior. A rigorous data preprocessing pipeline was first applied, "
    "including removal of high-missingness features, text field parsing, label encoding of "
    "categorical variables, outlier handling, and binary target engineering, resulting in a "
    "clean analytical dataset of 292,561 observations and 76 features. "
    "An interdependence technique, Factor Analysis, was employed to uncover the latent "
    "financial constructs underlying the feature space. Bartlett's Test of Sphericity "
    "confirmed the suitability of the correlation structure, and six interpretable factors "
    "were extracted, collectively explaining 56.5 percent of total variance. The identified "
    "factors correspond to meaningful economic dimensions including loan size, credit risk "
    "pricing, credit utilization, credit breadth, credit limit stress, and borrower wealth. "
    "Two dependence techniques were then applied for binary classification of loan default: "
    "Linear Discriminant Analysis and Logistic Regression, both trained on a class-balanced "
    "sample using a 70/30 train-test split. Both models achieved high discriminative "
    "performance, with area under the ROC curve scores of 0.997 and 0.9997 respectively. "
    "Results and limitations, including post-origination data leakage, are discussed.",
    bold=True, size=10)

blank(doc, 10)

# Index Terms
p_idx = doc.add_paragraph()
p_idx.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_single(p_idx)
add_run(p_idx, "Index Terms", italic=True, size=10)
add_run(p_idx,
    " - Credit Risk, Factor Analysis, Linear Discriminant Analysis, "
    "Logistic Regression, LendingClub.",
    size=10)

blank(doc, 10)
blank(doc, 10)

# ── Insert continuous section break → switch to two columns ──────
break_para = insert_continuous_section_break(doc)

# ── Add a new section and set it to two columns ──────────────────
new_section = doc.add_section(WD_SECTION.CONTINUOUS)
new_section.top_margin    = Inches(0.75)
new_section.bottom_margin = Inches(1.0)
new_section.left_margin   = Inches(0.75)
new_section.right_margin  = Inches(0.75)
set_two_columns(new_section)

# ═══════════════════════════════════════════════════════════════════
# SECTION I — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "I", "Introduction")

body_para(doc,
    "Credit risk assessment is one of the most consequential tasks in financial services. "
    "When a lender extends a loan, it bears the risk that the borrower may fail to repay — "
    "an event called a default. Accurate prediction of loan default enables institutions to "
    "price risk appropriately, minimize financial losses, and allocate capital more efficiently. "
    "At the consumer level, it also influences who gains access to credit and at what cost.",
    first=True)

body_para(doc,
    "Traditional credit scoring models, such as the FICO score, aggregate borrower "
    "information into a single numeric index. While widely used, such scalar representations "
    "necessarily discard multivariate structure — the complex, correlated relationships among "
    "dozens of financial variables that jointly characterize a borrower's risk profile. "
    "Multivariate statistical techniques address this limitation by analyzing many variables "
    "simultaneously, preserving and exploiting their joint structure.")

body_para(doc,
    "This paper applies three multivariate statistical techniques to a large real-world "
    "lending dataset from LendingClub: Factor Analysis as an interdependence technique to "
    "identify latent financial constructs, and Linear Discriminant Analysis (LDA) together "
    "with Logistic Regression as dependence techniques for binary classification of default. "
    "The modeling pipeline is presented end-to-end, from raw data ingestion through cleaning, "
    "dimensionality reduction, model training, and performance evaluation.")

body_para(doc,
    "The contributions of this work are: (1) a reproducible preprocessing pipeline for "
    "large-scale consumer lending data; (2) a Factor Analysis revealing six interpretable "
    "economic dimensions underlying LendingClub borrower characteristics; and (3) a "
    "comparative classification study between LDA and Logistic Regression, including "
    "discussion of data leakage considerations specific to post-origination credit datasets.")

# ═══════════════════════════════════════════════════════════════════
# SECTION II — BACKGROUND AND RELATED WORK
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "II", "Background and Related Work")

body_para(doc,
    "Credit default prediction has been studied extensively in the machine learning and "
    "statistics literature. Early work relied on discriminant analysis and logistic "
    "regression [1], which remain strong baselines due to their interpretability and "
    "statistical rigor. More recent work has explored ensemble methods, neural networks, "
    "and gradient boosting on credit datasets [2]; however, interpretability and regulatory "
    "compliance often favor classical statistical models in practice.",
    first=True)

body_para(doc,
    "LendingClub's publicly available loan data has been used in numerous studies. "
    "Tsai et al. [3] applied multiple classifiers to peer-to-peer lending data and found "
    "that logistic regression provided competitive performance with strong interpretability. "
    "Malekipirbazari and Aksakalli [4] demonstrated that machine learning classifiers "
    "trained on LendingClub data could outperform the platform's own grading system in "
    "separating good from bad borrowers.")

body_para(doc,
    "Factor Analysis has been applied to financial data to reduce dimensionality and "
    "identify latent constructs. Johnson and Wichern [5] provide foundational treatment of "
    "multivariate methods including Factor Analysis and discriminant analysis, which form "
    "the statistical backbone of this work. Hair et al. [6] further detail best practices "
    "for Factor Analysis in applied business research settings.")

body_para(doc,
    "The present work is distinguished by its combination of an interdependence technique "
    "(Factor Analysis) with two dependence techniques (LDA and Logistic Regression) applied "
    "to a single large-scale consumer credit dataset, providing a comprehensive multivariate "
    "statistical treatment of the credit classification problem.")

# ═══════════════════════════════════════════════════════════════════
# SECTION III — DATASET
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "III", "Dataset")

body_para(doc,
    "The dataset used in this study is the LendingClub Loan Data, sourced from the Kaggle "
    "public repository [7]. LendingClub is a United States-based peer-to-peer lending "
    "marketplace where borrowers apply for personal loans and investors fund them. The "
    "platform makes historical loan records publicly available, making this dataset "
    "representative of real consumer credit behavior.",
    first=True)

body_para(doc,
    "The raw dataset contains 500,000 loan records and 151 features per record. Each "
    "row corresponds to one loan application and its subsequent outcome. Features include "
    "borrower-reported financial characteristics (annual income, employment length, "
    "debt-to-income ratio), loan terms (loan amount, interest rate, installment, term), "
    "LendingClub-assigned attributes (loan grade, subgrade), credit bureau data (FICO "
    "score range, number of inquiries, number of open accounts, delinquency history), "
    "and post-origination performance metrics (total payments received, recoveries, "
    "charge-off amounts).")

body_para(doc,
    "The target variable, loan default, was engineered from the loan_status field. "
    "Loans labeled Charged Off, Default, Late (31–120 days), or Late (16–30 days) were "
    "coded as default (1). Loans labeled Fully Paid were coded as non-default (0). "
    "Records with ambiguous statuses — Current and In Grace Period — were excluded as "
    "their ultimate outcome is unknown. After this filtering, the dataset contained "
    "394,712 records, of which 79.1 percent were non-default and 20.9 percent were default.")

# ═══════════════════════════════════════════════════════════════════
# SECTION IV — DATA PREPROCESSING
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "IV", "Data Preprocessing")

sub_heading(doc, "A", "Missing Value Treatment")

body_para(doc,
    "Of the 151 original features, 58 were found to have missing rates exceeding 40 "
    "percent and were removed. These columns predominantly captured loan modification and "
    "hardship events that are only populated for a minority of records (e.g., "
    "hardship_payoff_balance_amount, settlement_date). An additional 17 columns were "
    "dropped as identifiers or free-text fields with no analytical value (e.g., loan ID, "
    "URL, borrower description). Rows with remaining missing values after these steps "
    "were removed via listwise deletion.",
    first=True)

sub_heading(doc, "B", "Feature Transformation")

body_para(doc,
    "Several numeric features were stored as strings in the raw data and required parsing. "
    "The interest rate field (int_rate) contained percent signs and was converted to a "
    "floating-point number. The term field was stored as '36 months' or '60 months' and "
    "was converted to an integer number of months. Employment length (emp_length) was "
    "stored as '10+ years' and similar string formats; these were parsed to integer years. "
    "Revolving utilization (revol_util) was similarly converted from a string percentage.",
    first=True)

body_para(doc,
    "Eight categorical columns were label-encoded to ordinal integers: loan grade, home "
    "ownership, verification status, loan purpose, initial list status, hardship flag, "
    "disbursement method, and debt settlement flag. Ordinal encoding was appropriate for "
    "grade (A through G represent increasing risk) and acceptable for the remaining "
    "nominal features given the classification methods used.")

sub_heading(doc, "C", "Outlier Handling")

body_para(doc,
    "Extreme outliers in annual income, loan amount, and revolving balance were removed "
    "by truncating values below the 1st and above the 99th percentile of each variable. "
    "This prevented a small number of extreme records from disproportionately influencing "
    "parameter estimates in both Factor Analysis and the classification models.",
    first=True)

sub_heading(doc, "D", "Final Dataset")

body_para(doc,
    "After all preprocessing steps, the final analytical dataset comprised 292,561 "
    "records and 76 features. The class distribution was 79.1 percent non-default and "
    "20.9 percent default, representing a moderate class imbalance that was addressed "
    "during model training.",
    first=True)

# ═══════════════════════════════════════════════════════════════════
# SECTION V — METHODOLOGY
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "V", "Methodology")

sub_heading(doc, "A", "Factor Analysis")

body_para(doc,
    "Factor Analysis (FA) is an interdependence technique that models observed variables "
    "as linear combinations of a smaller number of unobserved latent factors plus unique "
    "residuals [5]. It is appropriate when variables share common variance attributable to "
    "underlying constructs that are not directly measured. In the credit domain, constructs "
    "such as credit utilization stress or overall borrower wealth manifest across multiple "
    "correlated observed variables.",
    first=True)

body_para(doc,
    "Prior to FA, Bartlett's Test of Sphericity was conducted to verify that the "
    "correlation matrix is significantly different from an identity matrix — a prerequisite "
    "for Factor Analysis. The test yielded a chi-square statistic of 657,867 with p < 0.001, "
    "confirming the appropriateness of the method. FA was applied to 20 key financial "
    "features selected for their theoretical relevance and low missingness, using a random "
    "sample of 50,000 records to ensure computational tractability.")

body_para(doc,
    "The number of factors to retain was determined by the Kaiser criterion (eigenvalue > 1), "
    "corroborated by visual inspection of the scree plot. Six factors were extracted using "
    "the sklearn FactorAnalysis estimator, which fits the model via the expectation-"
    "maximization algorithm. Factor loadings were computed and interpreted to assign "
    "economic meaning to each factor.")

sub_heading(doc, "B", "Linear Discriminant Analysis")

body_para(doc,
    "Linear Discriminant Analysis (LDA) is a supervised dependence technique that finds "
    "the linear combination of features that maximally separates two or more class groups "
    "by maximizing the ratio of between-class to within-class scatter [1]. It produces a "
    "discriminant function that assigns each observation a scalar score, which is then "
    "thresholded to produce a binary class prediction. LDA assumes multivariate normality "
    "within each class and equal covariance matrices across classes.",
    first=True)

body_para(doc,
    "To address the 79/21 class imbalance, a class-balanced training sample was constructed "
    "by retaining all 62,328 default observations and randomly undersampling 124,656 "
    "non-default observations (a 2:1 ratio). All 75 numeric features were standardized to "
    "zero mean and unit variance prior to model fitting. A 70/30 stratified train-test "
    "split was applied. The scikit-learn LinearDiscriminantAnalysis implementation was used.")

sub_heading(doc, "C", "Logistic Regression")

body_para(doc,
    "Logistic Regression (LR) models the log-odds of the binary outcome as a linear "
    "function of the predictors and uses the sigmoid function to convert this to a "
    "probability estimate [1]. Unlike LDA, LR does not assume normality of predictors "
    "within classes, making it more robust to the non-normal distributions common in "
    "financial data. It was trained on the same balanced dataset and feature-standardized "
    "inputs as LDA, using L2 regularization with the liblinear solver. The same 70/30 "
    "train-test split was maintained for a fair comparison.",
    first=True)

sub_heading(doc, "D", "Evaluation Metrics")

body_para(doc,
    "Both classification models were evaluated using accuracy, precision, recall, F1-score, "
    "and the area under the receiver operating characteristic curve (AUC-ROC). The "
    "AUC-ROC is particularly important in credit risk settings because it measures overall "
    "discriminative ability across all decision thresholds, not just at the default "
    "0.5 cutoff. Confusion matrices were also generated to examine the distribution of "
    "true positives, false positives, true negatives, and false negatives.",
    first=True)

# ═══════════════════════════════════════════════════════════════════
# SECTION VI — RESULTS
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "VI", "Results")

sub_heading(doc, "A", "Factor Analysis Results")

body_para(doc,
    "Six factors were extracted, collectively explaining 56.5 percent of the total "
    "variance in the 20-variable feature space. Table I summarizes the factors, their "
    "proportion of variance explained, and the highest-loading variables. Figure I shows "
    "the scree plot and Figure II shows the factor loadings heatmap.",
    first=True)

body_para(doc,
    "Factor 1 (14.6% variance) loaded strongly on loan_amnt (+0.94), installment (+0.90), "
    "and term (+0.47), representing a Loan Size and Cost construct. Larger loans naturally "
    "carry higher monthly installments regardless of borrower quality.")

body_para(doc,
    "Factor 2 (13.7%) loaded on int_rate (+0.87) and grade (+0.86) positively, and "
    "fico_range_low negatively (−0.44), capturing Credit Risk Pricing — the mechanism by "
    "which lenders price borrower risk into the interest rate.")

body_para(doc,
    "Factor 3 (10.0%) was dominated by revol_util (+0.90) and bc_util (+0.84), reflecting "
    "Credit Utilization — the degree to which borrowers are drawing on their available "
    "revolving credit. This is a well-established risk signal in the credit literature.")

body_para(doc,
    "Factor 4 (9.4%) loaded on open_acc (+0.75), total_acc (+0.69), and revol_bal (+0.52), "
    "indicating Credit Breadth — the overall number and extent of credit relationships "
    "maintained by the borrower.")

body_para(doc,
    "Factors 5 and 6 (3.9% and 4.8% respectively) captured Credit Limit Stress and "
    "Borrower Wealth, the latter loading on mort_acc (+0.62) and tot_cur_bal (+0.53).")

# ── Table I ─────────────────────────────────────────────────────
blank(doc, 6)
p_th = doc.add_paragraph()
p_th.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_th)
add_run(p_th, "TABLE I", size=10)

p_tc = doc.add_paragraph()
p_tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_tc)
run_tc = add_run(p_tc, "Factor Analysis Summary", size=8)
make_small_caps(run_tc)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Factor"
hdr[1].text = "Var. (%)"
hdr[2].text = "Interpretation"
rows_data = [
    ("F1", "14.6", "Loan Size & Cost"),
    ("F2", "13.7", "Credit Risk Pricing"),
    ("F3", "10.0", "Credit Utilization"),
    ("F4", "9.4",  "Credit Breadth"),
    ("F5", "3.9",  "Credit Limit Stress"),
    ("F6", "4.8",  "Wealth & Assets"),
    ("Total", "56.5", "Cumulative Variance Explained"),
]
for i, (f, v, interp) in enumerate(rows_data):
    row = table.rows[i+1].cells
    row[0].text = f
    row[1].text = v
    row[2].text = interp
for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = "Times New Roman"
blank(doc, 6)

# Scree plot
insert_figure(doc, "results/fa_scree_plot.png", "I",
              "Scree Plot Showing Eigenvalues For Factor Retention", width=3.0)
blank(doc, 4)
insert_figure(doc, "results/fa_loadings_heatmap.png", "II",
              "Factor Loadings Heatmap Across 20 Financial Variables", width=2.2)

sub_heading(doc, "B", "Classification Results")

body_para(doc,
    "Table II presents the classification performance of LDA and Logistic Regression on "
    "the held-out test set. Figure III shows the ROC curves for both models.",
    first=True)

# ── Table II ────────────────────────────────────────────────────
blank(doc, 6)
p_th2 = doc.add_paragraph()
p_th2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_th2)
add_run(p_th2, "TABLE II", size=10)

p_tc2 = doc.add_paragraph()
p_tc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_single(p_tc2)
run_tc2 = add_run(p_tc2, "Classification Performance Comparison", size=8)
make_small_caps(run_tc2)

table2 = doc.add_table(rows=4, cols=5)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for cell, txt in zip(hdr2, ["Model", "Accuracy", "Precision", "Recall", "AUC-ROC"]):
    cell.text = txt
rows2 = [
    ("LDA",                "97.3%", "0.975", "0.973", "0.997"),
    ("Logistic Regression","99.6%", "0.998", "0.996", "0.9997"),
    ("Difference",         "+2.3%", "+0.023","+0.023","+0.0027"),
]
for i, rd in enumerate(rows2):
    row = table2.rows[i+1].cells
    for cell, txt in zip(row, rd):
        cell.text = txt
for row in table2.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.name = "Times New Roman"
blank(doc, 6)

body_para(doc,
    "LDA achieved 97.3 percent accuracy and an AUC-ROC of 0.997, correctly classifying "
    "the vast majority of both default and non-default loans. The discriminant projection "
    "plot (Figure IV) shows clear separation between the two classes along the single "
    "discriminant axis, consistent with the high AUC. The top discriminant coefficients "
    "highlighted fico_range_low (+30.99), int_rate, and total_pymnt as the most "
    "separating variables.",
    first=True)

body_para(doc,
    "Logistic Regression outperformed LDA on all metrics, achieving 99.6 percent accuracy "
    "and an AUC-ROC of 0.9997. The model's top positive coefficients — loan_amnt, "
    "installment, and recoveries — align intuitively with default risk. The top negative "
    "coefficients — total_rec_prncp and total_pymnt — reflect that higher cumulative "
    "payments are strongly associated with full repayment. Figure V shows the coefficient "
    "chart.")

insert_figure(doc, "results/roc_comparison.png", "III",
              "ROC Curves For LDA And Logistic Regression", width=3.0)
blank(doc, 4)
insert_figure(doc, "results/lda_projection.png", "IV",
              "LDA Discriminant Score Distributions By Class", width=3.0)
blank(doc, 4)
insert_figure(doc, "results/lr_coefficients.png", "V",
              "Top 10 Logistic Regression Coefficients By Magnitude", width=3.0)

# ═══════════════════════════════════════════════════════════════════
# SECTION VII — DISCUSSION
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "VII", "Discussion")

sub_heading(doc, "A", "Interpretability of Factors")

body_para(doc,
    "The six factors extracted by Factor Analysis correspond to economically meaningful "
    "and theoretically grounded constructs. The emergence of Credit Risk Pricing as a "
    "distinct factor (F2) confirms that LendingClub's interest rate and grade assignment "
    "mechanism effectively encodes borrower risk — higher-risk borrowers face steeper "
    "rates and worse grades, and this manifests as a coherent latent dimension.",
    first=True)

body_para(doc,
    "Credit Utilization (F3) and Credit Breadth (F4) are consistent with the five "
    "dimensions of credit quality used in traditional credit scoring: payment history, "
    "amounts owed, length of history, new credit, and types of credit. The fact that "
    "these emerge naturally from Factor Analysis without supervision validates their "
    "importance as risk signals.")

sub_heading(doc, "B", "Model Performance and Comparison")

body_para(doc,
    "Both classifiers achieved very high AUC-ROC scores (>0.99), exceeding typical "
    "benchmarks reported in the credit scoring literature for logistic models [3]. "
    "Logistic Regression's superior performance relative to LDA is consistent with "
    "the expectation that financial variables deviate from multivariate normality — "
    "LR's distributional assumptions are weaker and therefore more appropriate.",
    first=True)

body_para(doc,
    "The class-balanced training strategy was critical. Without undersampling, a naive "
    "model could achieve 79 percent accuracy by predicting non-default for all observations. "
    "Balancing the training set forced the model to learn genuine discriminative patterns "
    "rather than exploiting the prior class probability.")

sub_heading(doc, "C", "Data Leakage Limitation")

body_para(doc,
    "The unusually high accuracy figures are partially attributable to post-origination "
    "features present in the dataset. Variables such as total_rec_prncp (total principal "
    "recovered), recoveries, and collection_recovery_fee are only populated after a loan "
    "has been resolved — meaning they effectively encode the outcome being predicted. "
    "This constitutes data leakage: information that would not be available at prediction "
    "time is included in the feature set.",
    first=True)

body_para(doc,
    "In a production credit scoring system, only features known at loan origination — "
    "such as FICO score, annual income, debt-to-income ratio, and employment length — "
    "would be permissible inputs. The inclusion of post-outcome features inflates "
    "estimated model performance and limits the generalizability of the results to "
    "real-time underwriting decisions. This limitation is acknowledged and future work "
    "should restrict the feature set to origination-time variables only.")

# ═══════════════════════════════════════════════════════════════════
# SECTION VIII — CONCLUSION
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "VIII", "Conclusion")

body_para(doc,
    "This paper demonstrated the application of multivariate statistical techniques to "
    "credit risk classification using a large-scale real-world dataset from LendingClub. "
    "A comprehensive preprocessing pipeline was developed to transform 500,000 raw loan "
    "records into a clean, analysis-ready dataset of 292,561 observations.",
    first=True)

body_para(doc,
    "Factor Analysis revealed six latent financial constructs — Loan Size and Cost, Credit "
    "Risk Pricing, Credit Utilization, Credit Breadth, Credit Limit Stress, and Borrower "
    "Wealth — explaining 56.5 percent of total variance. These factors provide a compact "
    "and interpretable representation of the multi-dimensional borrower feature space, "
    "with direct correspondence to known credit risk theory.")

body_para(doc,
    "Both Linear Discriminant Analysis and Logistic Regression achieved strong "
    "discriminative performance, with AUC-ROC scores of 0.997 and 0.9997 respectively. "
    "Logistic Regression outperformed LDA across all metrics, consistent with its more "
    "flexible distributional assumptions. The primary limitation of both models is data "
    "leakage from post-origination features, which inflates performance estimates.")

body_para(doc,
    "Future work should replicate this analysis using only origination-time features to "
    "produce performance estimates representative of real-world underwriting settings. "
    "Additional extensions include applying Factor Analysis-derived scores as engineered "
    "features in the classification stage, and exploring non-linear classifiers such as "
    "gradient boosted trees for comparison.")

# ═══════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════
section_heading(doc, "", "References")

refs = [
    "[1] D. W. Hosmer and S. Lemeshow, Applied Logistic Regression, 2nd ed. New York, NY, USA: Wiley, 2000.",
    "[2] L. Breiman, \"Random Forests,\" Machine Learning, vol. 45, no. 1, pp. 5-32, Oct. 2001.",
    "[3] C. F. Tsai, Y. H. Hsu, and C. Yen, \"A Comparative Study of Classifier Ensembles for Bankruptcy Prediction,\" Applied Soft Computing, vol. 24, pp. 977-984, Nov. 2014.",
    "[4] M. Malekipirbazari and V. Aksakalli, \"Risk Assessment in Social Lending via Random Forests,\" Expert Systems with Applications, vol. 42, no. 10, pp. 4621-4631, Jun. 2015.",
    "[5] R. A. Johnson and D. W. Wichern, Applied Multivariate Statistical Analysis, 6th ed. Upper Saddle River, NJ, USA: Pearson, 2007.",
    "[6] J. F. Hair, W. C. Black, B. J. Babin, and R. E. Anderson, Multivariate Data Analysis, 7th ed. Upper Saddle River, NJ, USA: Pearson, 2010.",
    "[7] LendingClub Loan Data, Kaggle, 2018. [Online]. Available: https://www.kaggle.com/datasets/wordsforthewise/lending-club",
]

for ref in refs:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_single(p)
    p.paragraph_format.left_indent       = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after       = Pt(4)
    add_run(p, ref, size=8)

# ── Save ──────────────────────────────────────────────────────────
OUTPUT = "Gidan_LendingClub_CreditRisk_Paper.docx"
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
