"""
Generate ISEC-2017-formatted Word document for the S&P 500 Outperformer Prediction paper.
Authors: Lana Gidan, Matthew Golubow, Shahd Tarman — Binghamton University
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import os

VIZ_DIR = 'viz'
OUTPUT  = 'SP500_Technical_Report_LanaGidan.docx'

# ── Low-level helpers ──────────────────────────────────────────────────────────

def make_el(tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    return el


def font(run, size=10, bold=False, italic=False,
         small_caps=False, all_caps=False, underline=False):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.underline = underline
    if small_caps: run.font.small_caps = True
    if all_caps:   run.font.all_caps   = True


def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             sb=0, sa=0, first_indent=None, left_indent=None):
    pf = p.paragraph_format
    pf.alignment          = align
    pf.space_before       = Pt(sb)
    pf.space_after        = Pt(sa)
    pf.line_spacing_rule  = WD_LINE_SPACING.SINGLE
    if first_indent is not None:
        pf.first_line_indent = Inches(first_indent)
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)


def inject_2col_section_break(para):
    """
    Attach a continuous section break to `para` and switch
    the PREVIOUS section to 1 column (default), so the new
    section (body) becomes 2-column.
    """
    pPr = para._p.get_or_add_pPr()
    sectPr = make_el('w:sectPr')
    # section type = continuous
    t = make_el('w:type'); t.set(qn('w:val'), 'continuous')
    sectPr.append(t)
    # 2 columns, 0.20 inch = 288 twips spacing
    cols = make_el('w:cols')
    cols.set(qn('w:num'),        '2')
    cols.set(qn('w:space'),      '288')
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)
    pPr.append(sectPr)


def set_doc_2col(doc):
    """Set the final/main section of the document to 2 columns."""
    sectPr = doc.sections[-1]._sectPr
    # remove any existing cols
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = make_el('w:cols')
    cols.set(qn('w:num'),        '2')
    cols.set(qn('w:space'),      '288')
    cols.set(qn('w:equalWidth'), '1')
    sectPr.append(cols)


def remove_para_border(p):
    pPr = p._p.get_or_add_pPr()
    pb  = make_el('w:pBdr')
    for side in ('top','left','bottom','right'):
        el = make_el(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        pb.append(el)
    pPr.append(pb)


# ── Content helpers ────────────────────────────────────────────────────────────

def body_para(doc, text='', first=False):
    p = doc.add_paragraph()
    para_fmt(p, first_indent=(0 if first else 0.25))
    if text:
        r = p.add_run(text)
        font(r)
    return p


def section_heading(doc, roman, title):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=8)
    r = p.add_run(f'{roman}. {title}')
    font(r, bold=True, small_caps=True)
    return p


def sub_heading(doc, letter, title):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, sb=6, sa=6)
    r = p.add_run(f'{letter}. {title}')
    font(r, italic=True)
    return p


def add_figure(doc, roman, filename, caption):
    img_path = os.path.join(VIZ_DIR, filename)
    # image paragraph
    p_img = doc.add_paragraph()
    para_fmt(p_img, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=2)
    p_img.add_run().add_picture(img_path, width=Inches(3.2))
    # FIGURE N heading
    p_hd = doc.add_paragraph()
    para_fmt(p_hd, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=1)
    r = p_hd.add_run(f'FIGURE {roman}')
    font(r, size=10, all_caps=True)
    # caption
    p_cap = doc.add_paragraph()
    para_fmt(p_cap, align=WD_ALIGN_PARAGRAPH.CENTER, sb=1, sa=6)
    r = p_cap.add_run(caption)
    font(r, size=8, small_caps=True)
    return p_img


def add_table_heading(doc, roman, caption):
    p_hd = doc.add_paragraph()
    para_fmt(p_hd, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=1)
    r = p_hd.add_run(f'TABLE {roman}')
    font(r, size=10, all_caps=True)
    p_cap = doc.add_paragraph()
    para_fmt(p_cap, align=WD_ALIGN_PARAGRAPH.CENTER, sb=1, sa=3)
    r = p_cap.add_run(caption)
    font(r, size=8, small_caps=True)


def tbl_cell(cell, text, sz=8, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
             italic=False, small_caps=False):
    cell.text = ''
    p = cell.paragraphs[0]
    para_fmt(p, align=align, sb=1, sa=1)
    r = p.add_run(text)
    font(r, size=sz, bold=bold, italic=italic, small_caps=small_caps)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        tbl_cell(tbl.rows[0].cells[i], h, bold=True)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl_cell(tbl.rows[ri+1].cells[ci], val)
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Inches(w)
    return tbl


def spacer(doc, size=10):
    p = doc.add_paragraph()
    para_fmt(p, sb=0, sa=0)
    r = p.add_run('')
    font(r, size=size)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page layout ───────────────────────────────────────────────────────────────
sec0 = doc.sections[0]
sec0.page_width       = Inches(8.5)
sec0.page_height      = Inches(11)
sec0.top_margin       = Inches(0.75)
sec0.bottom_margin    = Inches(1.0)
sec0.left_margin      = Inches(0.75)
sec0.right_margin     = Inches(0.75)
sec0.header_distance  = Inches(0.5)
sec0.footer_distance  = Inches(0.5)

# Default style
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(10)

# ── TITLE (24 pt, centered, not bold) ────────────────────────────────────────
p_title = doc.add_paragraph()
para_fmt(p_title, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)
r = p_title.add_run('Using Logistic Regression to Determine U.S. Equity Performance')
font(r, size=24)

# 11 pt blank line
p_blank = doc.add_paragraph()
para_fmt(p_blank, sb=0, sa=0)
r = p_blank.add_run('')
font(r, size=11)

# ── AUTHORS (11 pt, centered) ─────────────────────────────────────────────────
p_auth = doc.add_paragraph()
para_fmt(p_auth, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)
r = p_auth.add_run('Lana Gidan, Matthew Golubow, and Shahd Tarman')
font(r, size=11)

# ── AFFILIATION (10 pt, centered) ─────────────────────────────────────────────
p_aff = doc.add_paragraph()
para_fmt(p_aff, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)
r = p_aff.add_run('Binghamton University, lgidan@binghamton.edu, mgolubo1@binghamton.edu, starman@binghamton.edu')
font(r, size=10)

# 2 x 10 pt blank lines
for _ in range(2):
    spacer(doc, 10)

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
p_abs = doc.add_paragraph()
para_fmt(p_abs, sb=0, sa=0)
r_label = p_abs.add_run('Abstract\u2014')
font(r_label, bold=True, italic=True)
abstract_text = (
    'Predicting which stocks will outperform a broad market index is a central problem '
    'in quantitative finance. This paper applies logistic regression to quarterly financial '
    'statement data for S&P 500 constituent companies over a 15-year training window '
    '(Q1 2010\u2013Q4 2024) to predict whether each company will outperform the S&P 500 '
    'index in the following quarter. The methodology follows the benchmark established by '
    'Ananthakumar and Sarkar (2017), including Variance Inflation Factor (VIF) filtering '
    'at a threshold of 2.5 and Principal Component Analysis (PCA) with a per-component '
    'variance threshold of 5%, reducing 24 engineered features to 9 uncorrelated principal '
    'components. A global pooled logistic regression model achieves 53.2% accuracy and an '
    'AUC of 0.5381, which is statistically significant (p < 0.0001) but falls below the '
    'benchmark accuracy of 71.2%. Segmentation experiments across 45 configurations\u2014using '
    'GICS sector labels, market capitalization tiers, and K-Means clustering\u2014demonstrate '
    'that sector-specific models substantially improve performance. The best configuration, '
    'Communication Services with K-Means clustering (k=2, Cluster 0), achieves 63.9% accuracy '
    'and AUC 0.6902. A Gradient Boosting benchmark (AUC 0.6589) confirms the presence of '
    'non-linear patterns beyond the reach of logistic regression. Out-of-sample validation '
    'on 1,219 stock-quarter predictions across all four quarters of 2025 yields accuracy '
    'ranging from 44.6% (Q1 2025, a bear-market quarter) to 57.5% (Q4 2025). A simulated '
    'long portfolio of predicted outperformers underperformed the S&P 500 in Q2 and Q3 2025 '
    'but marginally beat it in Q4 2025, consistent with the semi-strong form of the '
    'Efficient Market Hypothesis.'
)
r_body = p_abs.add_run(abstract_text)
font(r_body, bold=True)

# ── INDEX TERMS ───────────────────────────────────────────────────────────────
p_idx = doc.add_paragraph()
para_fmt(p_idx, sb=4, sa=0)
r_il = p_idx.add_run('Index Terms\u2014')
font(r_il, italic=True)
r_it = p_idx.add_run(
    'Efficient Market Hypothesis, Logistic Regression, '
    'Principal Component Analysis, S&P 500, Stock Outperformance Prediction.'
)
font(r_it)

# ── Continuous section break → 2-column body ──────────────────────────────────
p_brk = doc.add_paragraph()
para_fmt(p_brk, sb=10, sa=0)
inject_2col_section_break(p_brk)
set_doc_2col(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION I — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'I', 'Introduction')

body_para(doc, first=True, text=(
    'The Securities and Exchange Commission (SEC) mandates that all publicly traded companies '
    'disclose quarterly financial statements via Form 10-Q. These disclosures provide a rich, '
    'publicly available time series of financial ratios that analysts routinely use to evaluate '
    'company health and relative value.'
))
body_para(doc, text=(
    'This paper addresses a specific question: can these publicly disclosed quarterly financials '
    'predict, one quarter in advance, whether a given S&P 500 company\'s stock return will '
    'exceed that of the index itself? A positive answer would have direct implications for active '
    'portfolio management and would challenge the semi-strong form of the Efficient Market '
    'Hypothesis (EMH), which asserts that all publicly available information is already reflected '
    'in current prices.'
))
body_para(doc, text=(
    'The study is structured around three research questions: (Q1) Can logistic regression '
    'applied to quarterly financial ratios statistically discriminate outperformers from '
    'underperformers? (Q2) Does sector segmentation materially improve predictive accuracy over '
    'a global pooled model? (Q3) Can model-predicted outperformers be used to construct a '
    'portfolio that beats the S&P 500 on a quarterly basis?'
))
body_para(doc, text=(
    'We benchmark against Ananthakumar and Sarkar (2017) [1], who reported 71.2% accuracy on '
    'S&P 500 constituents using a similar logistic regression framework with VIF-based variable '
    'selection. Their threshold of 2.5 for VIF filtering and their forward-looking label '
    'construction are adopted directly in this study.'
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION II — RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'II', 'Related Work')

body_para(doc, first=True, text=(
    'Quantitative stock selection has been studied extensively since Fama and French (1992) [2], '
    'who documented that book-to-market and size factors explain cross-sectional variation in '
    'equity returns beyond the Capital Asset Pricing Model. Piotroski (2000) [3] demonstrated '
    'that a simple nine-variable accounting-based score could separate high- from low-return '
    'value stocks with high accuracy.'
))
body_para(doc, text=(
    'More directly relevant to this work, Ananthakumar and Sarkar (2017) [1] applied logistic '
    'regression with VIF variable selection (threshold 2.5) to quarterly Compustat data for '
    'S&P 500 constituents, achieving 71.2% prediction accuracy. Their framework serves as the '
    'methodological and performance benchmark for this study.'
))
body_para(doc, text=(
    'Krauss, Do, and Huck (2017) [4] demonstrated that Gradient Boosting and Deep Neural '
    'Networks substantially outperform linear models for S&P 500 return prediction, motivating '
    'our use of Gradient Boosting as a diagnostic non-linear benchmark alongside the primary '
    'logistic regression model.'
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION III — VARIABLES IDENTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'III', 'Variables Identification')

sub_heading(doc, 'A', 'Dependent Variable')
body_para(doc, first=True, text=(
    'The binary outcome variable, outperformer_next, is defined as 1 if the total return '
    'of stock i in quarter t+1 exceeds the S&P 500 index return over the same period, and '
    '0 otherwise. This forward-looking label is constructed by shifting the contemporaneous '
    'outperformance indicator back by one period using groupby(gvkey).shift(-1), ensuring '
    'that each observation\'s features belong to quarter t while the label belongs to quarter t+1. '
    'No future information is used in any feature calculation.'
))

sub_heading(doc, 'B', 'Independent Variables')
body_para(doc, first=True, text=(
    'Twenty-four features are engineered from raw Compustat fields, FRED macro data, and their '
    'quarter-over-quarter differences. These span six economic categories:'
))

# Bullet-style list
bullets = [
    ('Profitability:', 'ROA, ROE, Gross Margin, Operating Margin, Net Margin'),
    ('Efficiency:', 'Asset Turnover'),
    ('Liquidity:', 'Current Ratio'),
    ('Leverage:', 'Debt-to-Equity'),
    ('Growth:', 'Revenue Growth (QoQ), Net Income Growth (QoQ)'),
    ('Valuation:', 'P/E Ratio, Book-to-Market'),
    ('Macro:', 'GDP Growth (FRED), CPI Inflation (FRED)'),
    ('Momentum (\u0394):', 'Quarter-over-quarter change in each of the 10 ratio-based features above'),
]
for label, content in bullets:
    p = doc.add_paragraph(style='List Bullet')
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0,
             first_indent=0, left_indent=0.25)
    rb = p.add_run(label + ' ')
    font(rb, bold=True)
    rc = p.add_run(content)
    font(rc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION IV — METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'IV', 'Methodology')

sub_heading(doc, 'A', 'Overview')
body_para(doc, first=True, text=(
    'The analytical pipeline proceeds through five stages: (1) data collection and ratio '
    'computation, (2) preprocessing including winsorization and standardization, '
    '(3) multicollinearity reduction via iterative VIF filtering, (4) dimensionality reduction '
    'via PCA, and (5) logistic regression estimation across multiple segmentation configurations.'
))

sub_heading(doc, 'B', 'VIF Filtering')
body_para(doc, first=True, text=(
    'Variance Inflation Factor (VIF) quantifies the extent to which a predictor\'s variance '
    'is inflated by its linear relationship with other predictors. For predictor j, '
    'VIF_j = 1 / (1 - R_j^2), where R_j^2 is the coefficient of determination from regressing '
    'predictor j on all other predictors. Following Ananthakumar and Sarkar (2017), a threshold '
    'of 2.5 is applied\u2014stricter than the conventional value of 5.0\u2014using an iterative '
    'procedure: at each step, the predictor with the highest VIF above the threshold is removed '
    'and VIF values are recomputed until all remaining predictors satisfy VIF_j \u2264 2.5.'
))

sub_heading(doc, 'C', 'Principal Component Analysis')
body_para(doc, first=True, text=(
    'PCA is applied to the VIF-filtered feature set to eliminate residual correlation and reduce '
    'dimensionality. Components are retained if their individual explained variance ratio exceeds '
    '5%, defined as: n_comp = max(1, sum(lambda_k / sum(lambda_j) >= 0.05)). This threshold '
    'was recommended by the faculty advisor to retain components that individually contribute '
    'meaningful variance, rather than the conventional 80% cumulative rule.'
))

sub_heading(doc, 'D', 'Logistic Regression')
body_para(doc, first=True, text=(
    'The probability that company i will outperform the S&P 500 in quarter t+1, given its '
    'PCA score vector z_i,t, is modeled as: P(y = 1 | z) = 1 / (1 + exp(-(beta_0 + beta^T z))). '
    'Statistical significance of individual components is assessed via Wald z-tests; overall '
    'model significance via the Likelihood Ratio Test. Model fit is reported using McFadden\'s '
    'Pseudo-R\u00b2. The classification threshold tau is optimized over the range [0.30, 0.75] '
    'on training data to maximize F1 score.'
))

sub_heading(doc, 'E', 'Segmentation Configurations')
body_para(doc, first=True, text=(
    'Five configurations are evaluated, totaling 45 distinct models: (1) global pooled, '
    '(2) market capitalization tier (Large Cap >$10B, Mid Cap $2B-$10B, Small Cap <$2B), '
    '(3) GICS sector (11 sectors), (4) GICS sector with K-Means clustering (k=2, applied '
    'within each sector on PCA scores), and (5) GICS sector with market capitalization '
    'as an additional covariate.'
))

sub_heading(doc, 'F', 'Gradient Boosting Benchmark')
body_para(doc, first=True, text=(
    'A Gradient Boosting Classifier (200 estimators, max depth 3, learning rate 0.1) is '
    'trained on the same 9 PCA components as a non-linear diagnostic benchmark. Its AUC '
    'relative to logistic regression quantifies the amount of non-linear signal present '
    'in the data that logistic regression cannot capture.'
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION V — DATA DESCRIPTION & PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'V', 'Data Description & Preprocessing')

sub_heading(doc, 'A', 'Data Sources')
body_para(doc, first=True, text=(
    'Data are drawn from three primary sources accessed via the Wharton Research Data Service '
    '(WRDS): Compustat Fundamentals Quarterly (fundq table) for 21 raw financial statement '
    'fields, CRSP for quarterly stock returns and S&P 500 benchmark returns, and the Federal '
    'Reserve Economic Data (FRED) API for GDP growth and CPI inflation. S&P 500 constituent '
    'membership is tracked through the WRDS index history file to ensure only actual index '
    'members are included at each point in time. GICS sector classifications are joined via '
    'company identifier (gvkey). The raw Compustat extract covers Q1 2010 through Q4 2024 '
    'and contains 30,841 company-quarter observations across approximately 500 constituents.'
))

sub_heading(doc, 'B', 'Feature Correlation')
body_para(doc, first=True, text=(
    'Figure VIII presents the lower-triangle Pearson correlation matrix for the 12 base '
    'financial ratios before VIF filtering. Strong correlations (|r| >= 0.50) among the '
    'profitability ratios (ROA, ROE, Net Margin, Operating Margin) confirm the need for '
    'VIF-based filtering prior to model estimation.'
))
add_figure(doc, 'VIII', '08_correlation_heatmap.png',
           'Pearson Correlation Matrix of 12 Base Financial Ratios Before VIF Filtering')

sub_heading(doc, 'C', 'Preprocessing Pipeline')
body_para(doc, first=True, text=(
    'Three preprocessing steps are applied in sequence to the 21,285 initial company-quarter '
    'observations. First, winsorization clips each ratio at its 1st and 99th percentile values, '
    'computed on training data only and applied identically to all test observations to prevent '
    'data leakage. Second, standardization transforms each ratio to zero mean and unit standard '
    'deviation using training-set statistics, ensuring no predictor dominates due to scale '
    'differences. Third, rows containing any missing value after ratio computation are removed; '
    'the primary sources of missingness are zero-revenue observations and the first quarter of '
    'each company\'s history (no lag available for delta features). The cleaned training dataset '
    'contains 16,589 complete observations.'
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION VI — ANALYSIS & IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'VI', 'Analysis & Implementation')

sub_heading(doc, 'A', 'VIF Analysis Results')
body_para(doc, first=True, text=(
    'Iterative VIF filtering at threshold 2.5 removes four of the 24 candidate features: '
    'Return on Assets (initial VIF 4.52), Net Profit Margin (5.86), Delta ROA (5.11), and '
    'Delta Operating Margin (3.68). These variables exhibit high collinearity with the retained '
    'profitability and efficiency features. The remaining 20 features all satisfy VIF_j <= 2.5. '
    'Figure III displays all 24 initial VIF values, with retained features in green and removed '
    'features in red. Table I summarizes the key results.'
))

add_figure(doc, 'III', '03_vif_chart.png',
           'Initial VIF Values for All 24 Candidate Features. Red Bars Indicate Removed '
           'Features (VIF > 2.5); Green Bars Indicate Retained Features.')

add_table_heading(doc, 'I', 'VIF Analysis Summary (Selected Features)')
add_table(doc,
    headers=['Feature', 'Initial VIF', 'Decision'],
    rows=[
        ['Net Profit Margin',    '5.86', 'Removed'],
        ['\u0394ROA (QoQ)',       '5.11', 'Removed'],
        ['ROA',                  '4.52', 'Removed'],
        ['\u0394Net Margin',     '4.78', 'Removed'],
        ['\u0394Op. Margin',     '3.68', 'Removed'],
        ['ROE',                  '2.59', 'Retained'],
        ['\u0394Gross Margin',   '2.62', 'Retained'],
        ['Revenue Growth',       '2.49', 'Retained'],
        ['Debt-to-Equity',       '2.25', 'Retained'],
        ['Operating Margin',     '1.80', 'Retained'],
    ],
    col_widths=[1.5, 0.9, 0.8]
)
spacer(doc)

sub_heading(doc, 'B', 'PCA Analysis Results')
body_para(doc, first=True, text=(
    'PCA applied to the 20 retained features yields 9 components each explaining at least '
    '5% of total variance, collectively accounting for 70.7% of total variance. '
    'Figure II presents the scree plot with the 5% retention threshold marked. '
    'Table II reports the variance explained by each retained component.'
))

add_figure(doc, 'II', '02_pca_scree.png',
           'PCA Scree Plot. Green Bars Are Retained Components (>=5% Variance). '
           'Blue Line Shows Cumulative Variance Explained.')

add_table_heading(doc, 'II', 'PCA Component Variance and Dominant Themes (Global Model)')
add_table(doc,
    headers=['Component', 'Dominant Theme', 'Variance (%)'],
    rows=[
        ['PC1', 'Profitability (margins, ROE)', '12.7'],
        ['PC2', 'Leverage & Liquidity',          '10.3'],
        ['PC3', 'Revenue Momentum',               '9.3'],
        ['PC4', 'Valuation (P/E, B/M)',           '7.9'],
        ['PC5', 'NI & Margin Delta',              '7.3'],
        ['PC6', 'Asset Efficiency',               '6.3'],
        ['PC7', 'Macro Environment',              '5.8'],
        ['PC8', 'Equity Momentum',                '5.7'],
        ['PC9', 'Mixed Residual',                 '5.3'],
        ['Total', '', '70.7'],
    ],
    col_widths=[0.85, 1.75, 0.85]
)
spacer(doc)

sub_heading(doc, 'C', 'Global Model Estimation')
body_para(doc, first=True, text=(
    'The global logistic regression model is estimated on 16,589 training observations using '
    'all 9 PCA components with balanced class weights. Table III reports coefficient estimates, '
    'Wald z-statistics, and significance levels. Four components reach statistical significance '
    'at p < 0.05: PC1 (profitability, z = 4.955), PC3 (revenue momentum, z = -2.377), PC6 '
    '(asset efficiency, z = -5.114), and PC7 (macro environment, z = -3.015). McFadden\'s '
    'Pseudo-R\u00b2 = 0.0029. The overall model is statistically significant (LR test p < 0.0001).'
))

add_table_heading(doc, 'III', 'PCA Component Coefficients and Wald Test Results (Global Model)')
add_table(doc,
    headers=['Comp.', 'Coeff.', 'Odds Ratio', 'z-Stat', 'Sig.'],
    rows=[
        ['PC1', ' 0.0486', '1.050', ' 4.955', '***'],
        ['PC2', ' 0.0016', '1.002', ' 0.151', ''],
        ['PC3', '-0.0272', '0.973', '-2.377', '**'],
        ['PC4', '-0.0031', '0.997', '-0.251', ''],
        ['PC5', ' 0.0025', '1.003', ' 0.194', ''],
        ['PC6', '-0.0709', '0.932', '-5.114', '***'],
        ['PC7', '-0.0435', '0.957', '-3.015', '***'],
        ['PC8', '-0.0168', '0.983', '-1.152', ''],
        ['PC9', ' 0.0111', '1.011', ' 0.734', ''],
    ],
    col_widths=[0.6, 0.65, 0.85, 0.7, 0.55]
)
p_note = doc.add_paragraph()
para_fmt(p_note, align=WD_ALIGN_PARAGRAPH.LEFT, sb=1, sa=4)
r = p_note.add_run('*** p<0.01;  ** p<0.05')
font(r, size=7, italic=True)

sub_heading(doc, 'D', 'Confusion Matrix')
body_para(doc, first=True, text=(
    'Figure VI presents the confusion matrix for the global model at classification '
    'threshold 0.50. The model produces 5,766 true positives (TP), 3,062 true negatives (TN), '
    '5,055 false positives (FP), and 2,706 false negatives (FN), yielding sensitivity of 68.1% '
    'and specificity of 37.7%. The high sensitivity relative to low specificity indicates the '
    'model is better calibrated to identify outperformers than underperformers, consistent with '
    'the balanced class weights applied during training.'
))
add_figure(doc, 'VI', '06_confusion_matrix.png',
           'Confusion Matrix at Classification Threshold 0.50 '
           '(Global Model, 16,589 Training Observations)')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION VII — RESULTS & INTERPRETATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'VII', 'Results & Interpretation')

sub_heading(doc, 'A', 'Global Model Performance')
body_para(doc, first=True, text=(
    'The global logistic regression model achieves 53.2% overall accuracy and AUC of 0.5381. '
    'Table IV compares logistic regression to the Gradient Boosting benchmark and the '
    'Ananthakumar & Sarkar (2017) paper benchmark. Figure I presents the ROC curves for '
    'both models alongside the random baseline diagonal.'
))

add_figure(doc, 'I', '01_roc_curve.png',
           'ROC Curves for Logistic Regression (AUC=0.5381) and Gradient Boosting '
           '(AUC=0.6589). Gold Dot Marks the Chosen Classification Threshold (0.50).')

add_table_heading(doc, 'IV', 'Global Model Comparison: Logistic Regression vs. Gradient Boosting')
add_table(doc,
    headers=['Model', 'Accuracy', 'AUC'],
    rows=[
        ['Logistic Regression (primary)',  '53.2%', '0.5381'],
        ['Gradient Boosting (benchmark)',  '55.4%', '0.6589'],
        ['Ananthakumar & Sarkar (2017)',   '71.2%', '\u2014'],
    ],
    col_widths=[2.1, 0.9, 0.8]
)
spacer(doc)

body_para(doc, text=(
    'The 12.1 percentage-point AUC gap between logistic regression and Gradient Boosting '
    'confirms significant non-linear interaction effects among financial ratios. The gap '
    'between either model and the 71.2% benchmark is attributable to the stricter 5% '
    'per-component PCA threshold, which retains fewer components than a cumulative 80% '
    'rule would, and to the 15-year training window that spans multiple market regimes '
    'with conflicting ratio-return relationships.'
))

sub_heading(doc, 'B', 'Segmentation Results')
body_para(doc, first=True, text=(
    'Figure IV and Table V present accuracy and AUC across all 11 GICS sectors. '
    'Sector-specific models consistently outperform the global pooled model (53.2%), '
    'confirming that financial ratios carry different predictive content across industries. '
    'Energy (61.7%), Communication Services (59.5%), and Real Estate (59.5%) are the '
    'strongest-performing sectors.'
))

add_figure(doc, 'IV', '04_sector_results.png',
           'Test Accuracy and AUC by GICS Sector. Dashed Lines Mark the Global Baseline '
           '(Accuracy 53.2%, AUC 0.5381).')

add_table_heading(doc, 'V', 'Logistic Regression Results by GICS Sector')
add_table(doc,
    headers=['Sector', 'Accuracy', 'AUC'],
    rows=[
        ['Energy',                  '61.7%', '0.6036'],
        ['Communication Services',  '59.5%', '0.6059'],
        ['Real Estate',             '59.5%', '0.6010'],
        ['Consumer Staples',        '55.7%', '0.5510'],
        ['Financials',              '55.0%', '0.5350'],
        ['Industrials',             '54.5%', '0.5350'],
        ['Utilities',               '54.5%', '0.5420'],
        ['Health Care',             '54.3%', '0.5578'],
        ['Consumer Discretionary',  '54.1%', '0.5305'],
        ['Information Technology',  '53.2%', '0.5420'],
        ['Materials',               '53.2%', '0.5481'],
    ],
    col_widths=[2.0, 0.95, 0.85]
)
spacer(doc)

body_para(doc, text=(
    'The best single configuration across all 45 tested is Communication Services with '
    'K-Means clustering (k=2), Cluster 0: 63.9% accuracy and AUC 0.6902 on 122 observations. '
    'This cluster captures companies with subscription-driven revenue models whose quarterly '
    'ratio patterns exhibit strong intra-cluster homogeneity, making their outperformance '
    'more predictable from financial statement data alone.'
))

sub_heading(doc, 'C', '2025 Out-of-Sample Validation')
body_para(doc, first=True, text=(
    'The model trained through Q4 2024 is applied to all four quarters of 2025 on 1,219 '
    'stock-quarter predictions. Figure VII shows the quarter-by-quarter accuracy trend. '
    'The anomalous Q1 2025 accuracy of 44.6% coincides with a broad market correction '
    '(S&P 500 declined 4.6%), suggesting that extreme bear-market dynamics lie outside '
    'the distribution of the training data. Accuracy recovers to 55.3%-57.5% in subsequent '
    'quarters as market conditions normalize.'
))

add_figure(doc, 'VII', '07_accuracy_trend_2025.png',
           '2025 Out-of-Sample Prediction Accuracy by Quarter. Dashed Lines Indicate '
           'the Random Baseline (50%), Training Accuracy (53.2%), and Benchmark (71.2%).')

sub_heading(doc, 'D', 'Portfolio Simulation')
body_para(doc, first=True, text=(
    'A simulated portfolio is constructed each quarter by allocating proportionally by '
    'market capitalization across all stocks predicted to outperform. Figure V compares '
    'quarterly portfolio returns against SPY. Table VI reports quarterly performance statistics.'
))

add_figure(doc, 'V', '05_portfolio_vs_spy.png',
           'Simulated Model Portfolio Returns vs. S&P 500 (SPY) by Quarter, 2025. '
           'Bottom Panel Shows Alpha (Portfolio Return \u2212 SPY Return).')

add_table_heading(doc, 'VI', '2025 Simulated Portfolio vs. S&P 500 (SPY)')
add_table(doc,
    headers=['Quarter', 'N Stocks', 'Model Return', 'SPY Return', 'Alpha'],
    rows=[
        ['Q2 2025', '387', '+6.39%', '+10.57%', '\u22124.18%'],
        ['Q3 2025', '391', '+6.74%', '+7.79%',  '\u22121.06%'],
        ['Q4 2025', ' 72', '+3.15%', '+2.35%',  '+0.81%'],
    ],
    col_widths=[0.85, 0.75, 1.0, 1.0, 0.75]
)
spacer(doc)

body_para(doc, text=(
    'The portfolio underperforms SPY in two of three comparable quarters, consistent with '
    'the semi-strong form EMH: publicly disclosed financial ratios do not systematically '
    'support sustained risk-adjusted outperformance of a passive index investment. The '
    'marginal positive alpha in Q4 2025 (+0.81%) is within normal estimation error and '
    'should not be interpreted as evidence of systematic edge.'
))

sub_heading(doc, 'E', 'Interpretation Summary')
body_para(doc, first=True, text=(
    'The results collectively support three conclusions. First, quarterly financial ratios '
    'contain statistically significant but modest predictive signal for next-quarter '
    'outperformance (Q1 confirmed; practical effect is limited). Second, sector homogeneity '
    'is a meaningful moderator of predictive power: segmenting by GICS sector and further by '
    'within-sector financial profile clusters materially improves AUC (Q2 confirmed). '
    'Third, the model does not reliably generate positive alpha against a passive S&P 500 '
    'strategy (Q3 not confirmed), consistent with efficient markets in a competitive, '
    'information-rich environment.'
))

# ══════════════════════════════════════════════════════════════════════════════
# SECTION VIII — CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, 'VIII', 'Conclusion')

body_para(doc, first=True, text=(
    'This study demonstrates that logistic regression applied to quarterly financial ratios '
    'yields statistically significant but practically modest S&P 500 outperformance predictions '
    'in a global pooled setting. Sector segmentation and within-sector clustering substantially '
    'improve discrimination, with the Communication Services Cluster 0 configuration reaching '
    '63.9% accuracy (AUC 0.6902). The persistent gap between logistic regression and Gradient '
    'Boosting AUC scores (0.5381 vs. 0.6589) indicates that non-linear interaction effects in '
    'financial ratios remain unexploited by linear models.'
))
body_para(doc, text=(
    'Out-of-sample 2025 results confirm that the model generalizes reasonably to unseen data '
    'but does not generate consistent portfolio alpha, supporting a partial endorsement of the '
    'semi-strong Efficient Market Hypothesis. The sharp drop in Q1 2025 accuracy during a '
    'market correction underscores the importance of regime-aware modeling.'
))
body_para(doc, text=(
    'Future work should explore non-linear models (Random Forest, XGBoost, LSTM networks), '
    'longer lag windows to capture multi-quarter momentum, and alternative segmentation '
    'criteria such as analyst coverage density or earnings surprise history. Incorporating '
    'alternative data sources such as sentiment scores from 10-Q filings may further close '
    'the gap with the 71.2% benchmark accuracy.'
))

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENT
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '', 'Acknowledgment')

body_para(doc, first=True, text=(
    'The authors thank the SSIE 605 course faculty at Binghamton University for guidance '
    'on methodology selection and the VIF and PCA threshold decisions applied in this study. '
    'Data were accessed through the Wharton Research Data Service (WRDS) institutional '
    'subscription provided by Binghamton University.'
))

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, '', 'References')

refs = [
    ('U. Ananthakumar and D. Sarkar, "Application of logistic regression in assessing '
     'stock performance," Proc. IEEE ISEC, 2017.'),
    ('E. F. Fama and K. R. French, "The cross-section of expected stock returns," '
     'Journal of Finance, vol. 47, no. 2, pp. 427\u2013465, 1992.'),
    ('J. D. Piotroski, "Value investing: The use of historical financial statement '
     'information to separate winners from losers," Journal of Accounting Research, '
     'vol. 38, pp. 1\u201341, 2000.'),
    ('C. Krauss, X. A. Do, and N. Huck, "Deep neural networks, gradient-boosted trees, '
     'random forests: Statistical arbitrage on the S&P 500," European Journal of '
     'Operational Research, vol. 259, no. 2, pp. 689\u2013702, 2017.'),
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=0, sa=2,
             first_indent=-0.2, left_indent=0.2)
    r = p.add_run(f'[{i}] {ref}')
    font(r, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}  ({os.path.getsize(OUTPUT)/1024:.0f} KB)')
