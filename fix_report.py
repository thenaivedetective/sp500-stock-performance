"""
Fix all identified errors in the student's Word report.
"""
from docx import Document
from docx.oxml.ns import qn
import copy, re

SRC = 'attached_assets/UsingLogisticRegressionToDetermineEquityAlpha_1778416155693.docx'
OUT = 'UsingLogisticRegressionToDetermineEquityAlpha_FIXED.docx'

doc = Document(SRC)

# ─────────────────────────────────────────────────────────────
# Helper: replace text inside a paragraph preserving runs/style
# ─────────────────────────────────────────────────────────────
def replace_in_para(para, old, new):
    """Replace old→new across all runs in a paragraph."""
    full = para.text
    if old not in full:
        return False
    # Rebuild using the first run to hold new text, clear rest
    inline = para.runs
    # Collect all text into run[0], blank others
    for i, run in enumerate(inline):
        run.text = run.text.replace(old, new)
    return True

def replace_in_all_paras(old, new):
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            replace_in_para(para, old, new)
            count += 1
    return count

# ─────────────────────────────────────────────────────────────
# FIX 1 — Grammar: "may can still be" → "may still be"
# ─────────────────────────────────────────────────────────────
n = replace_in_all_paras('may can still be', 'may still be')
print(f'Fix 1 (grammar): {n} replacement(s)')

# ─────────────────────────────────────────────────────────────
# FIX 2 — Wrong figure ref: "Figure I presents the ROC" → "Figure V"
# ─────────────────────────────────────────────────────────────
n = replace_in_all_paras('Figure I presents the ROC curves', 'Figure V presents the ROC curves')
print(f'Fix 2 (Figure I→V ref): {n} replacement(s)')

# ─────────────────────────────────────────────────────────────
# FIX 3 — Inconsistent accuracy: "52.5%" → "53.2%" (only in testing section)
# ─────────────────────────────────────────────────────────────
for para in doc.paragraphs:
    if '52.5%' in para.text and 'in-sample' in para.text:
        replace_in_para(para, '52.5%', '53.2%')
        print('Fix 3 (52.5%→53.2%): done')
        break

# ─────────────────────────────────────────────────────────────
# FIX 4 — Best configuration claim: Energy → Comm Services Cluster 0
# ─────────────────────────────────────────────────────────────
OLD_BEST = ('The best single configuration across all models tested is the '
            'Energy sector model with 61.7% accuracy and .060 AUC. '
            'Splitting by market capitalization, as seen in Table VI, '
            'marginally increased accuracy against the global model, '
            'but not as much as the market segmentation did.')
NEW_BEST = ('The best single configuration across all 45 models tested is '
            'Communication Services Cluster 0, identified via K-Means clustering '
            'within the Communication Services sector, with AUC = 0.6902 and '
            '63.9% accuracy. Among the plain sector splits (without clustering), '
            'Energy leads with 61.7% accuracy and AUC = 0.6036. Splitting by '
            'market capitalization, as seen in Table VI, marginally increased '
            'accuracy against the global model, but not as much as sector '
            'segmentation did.')

fixed = False
for para in doc.paragraphs:
    if 'best single configuration' in para.text and 'Energy sector model' in para.text:
        # Replace text across all runs
        full_text = para.text
        # Clear existing runs and put corrected text in first run
        for i, run in enumerate(para.runs):
            if i == 0:
                run.text = NEW_BEST
            else:
                run.text = ''
        fixed = True
        print('Fix 4 (best config claim): done')
        break
if not fixed:
    print('Fix 4: paragraph not found — check manually')

# ─────────────────────────────────────────────────────────────
# FIX 5 — Portfolio figure label: second "FIGURE V" → "FIGURE VIII"
#         and the text reference "Figure V compares" → "Figure VIII compares"
# ─────────────────────────────────────────────────────────────
figv_count = 0
for para in doc.paragraphs:
    txt = para.text.strip()
    # The portfolio label appears AFTER the ROC curve label
    if txt == 'Figure V' or txt == 'FIGURE V':
        figv_count += 1
        if figv_count == 2:          # second occurrence = portfolio figure
            replace_in_para(para, 'Figure V', 'Figure VIII')
            replace_in_para(para, 'FIGURE V', 'FIGURE VIII')
            print('Fix 5a (portfolio figure label → FIGURE VIII): done')

# Fix text reference in portfolio paragraph
for para in doc.paragraphs:
    if 'Figure V compares' in para.text and 'portfolio' in para.text.lower():
        replace_in_para(para, 'Figure V compares', 'Figure VIII compares')
        print('Fix 5b (Figure V ref in portfolio text → VIII): done')
        break

# Also fix any caption line that says just "FIGURE V" for portfolio
# (sometimes it's a separate paragraph with just the caption text)
figv_count2 = 0
for para in doc.paragraphs:
    if 'FIGURE V' in para.text and 'Simulated' in para.text:
        replace_in_para(para, 'FIGURE V', 'FIGURE VIII')
        print('Fix 5c (Simulated portfolio figure → FIGURE VIII): done')

# ─────────────────────────────────────────────────────────────
# FIX 6 — Table 0 (VIF Table I): remove duplicate rows 9 and 12
#         Row 9: "Return on Equity" / 2.59 (duplicate of "ROE")
#         Row 12: "Operating Margin" / 1.80 (duplicate of row 5)
# ─────────────────────────────────────────────────────────────
vif_table = doc.tables[0]

def delete_row(table, row_index):
    tbl = table._tbl
    tr  = table.rows[row_index]._tr
    tbl.remove(tr)

# Delete in reverse order so indices don't shift
rows_to_delete = []
for ri, row in enumerate(vif_table.rows):
    cells = [c.text.strip() for c in row.cells]
    # Duplicate ROE row
    if cells[0] == 'Return on Equity' and cells[1] == '2.59':
        rows_to_delete.append(ri)
    # Second Operating Margin row (1.80)
    if cells[0] == 'Operating Margin' and cells[1] == '1.80':
        rows_to_delete.append(ri)

for ri in sorted(rows_to_delete, reverse=True):
    delete_row(vif_table, ri)
    print(f'Fix 6: deleted VIF table row {ri}')

# ─────────────────────────────────────────────────────────────
# FIX 7 — Table 4 (Market Cap Table VI): correct AUC values
#         Small Cap: 0.6036 → 0.5350
#         Large Cap: 0.6059 → 0.5578
#         Mid Cap:   0.6010 → 0.5420
# ─────────────────────────────────────────────────────────────
CORRECT_AUC = {
    'Small Cap': '0.5350',
    'Large Cap': '0.5578',
    'Mid Cap':   '0.5420',
}

mktcap_table = doc.tables[4]
for row in mktcap_table.rows:
    cap = row.cells[0].text.strip()
    if cap in CORRECT_AUC:
        auc_cell = row.cells[2]
        old_val = auc_cell.text.strip()
        for para in auc_cell.paragraphs:
            for run in para.runs:
                if old_val in run.text:
                    run.text = run.text.replace(old_val, CORRECT_AUC[cap])
        # If no runs found, set the paragraph text directly
        if auc_cell.text.strip() != CORRECT_AUC[cap]:
            for para in auc_cell.paragraphs:
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = CORRECT_AUC[cap]
                else:
                    para.add_run(CORRECT_AUC[cap])
        print(f'Fix 7: Market Cap {cap} AUC {old_val} → {CORRECT_AUC[cap]}')

# ─────────────────────────────────────────────────────────────
# FIX 8 — Table 5 (Portfolio): fill missing Alpha values
#         Q2 2025: −4.18%,  Q3 2025: −1.06%
# ─────────────────────────────────────────────────────────────
ALPHA_FIX = {
    'Q2 2025': '−4.18%',
    'Q3 2025': '−1.06%',
}

portfolio_table = doc.tables[5]
for row in portfolio_table.rows:
    qtr = row.cells[0].text.strip()
    if qtr in ALPHA_FIX:
        alpha_cell = row.cells[4]
        if not alpha_cell.text.strip():
            for para in alpha_cell.paragraphs:
                if para.runs:
                    para.runs[0].text = ALPHA_FIX[qtr]
                else:
                    para.add_run(ALPHA_FIX[qtr])
            print(f'Fix 8: {qtr} Alpha → {ALPHA_FIX[qtr]}')

# ─────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'\nSaved: {OUT}')

# Verification pass
doc2 = Document(OUT)
print('\n── Verification ──')
print('VIF table rows:', len(doc2.tables[0].rows), '(expected 11)')
for row in doc2.tables[4].rows:
    print(' MarketCap:', [c.text.strip() for c in row.cells])
for row in doc2.tables[5].rows:
    print(' Portfolio:', [c.text.strip() for c in row.cells])
